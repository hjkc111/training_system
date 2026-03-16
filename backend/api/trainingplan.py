from fastapi import APIRouter, HTTPException, Body, Response
from datetime import datetime, timedelta
import jwt
import os
import re
from io import BytesIO
# 终极兼容写法：同时支持Python3.8+所有版本，彻底解决导入报错
try:
    # Python3.9+ 优先用内置类型
    from typing import List, Optional, Dict, Any
except ImportError:
    # 兜底兼容
    from typing_extensions import List, Optional,   Dict, Any


# 第三方库
from docx import Document
import urllib.parse  # 新增：处理中文文件名编码

# 项目配置和工具
from config import FAKE_USERS
from models import LoginRequest, TrainingPlanGenerateRequest, TrainingPlanExportRequest
from utils.llm_utils import call_qwen_training_plan

# ------------------- 路由配置 -------------------
router = APIRouter(prefix="/api/trainingplan", tags=["训练日管理模块-训练计划"])

# ------------------- 常量定义 -------------------
# 文件名安全过滤正则（移除特殊字符）
SAFE_FILENAME_PATTERN = re.compile(r'[\\/:*?"<>|]')
# 默认导出文件名
DEFAULT_PLAN_FILENAME = "光电训练计划"


# ------------------- 生成光电训练计划接口 -------------------
@router.post("/generate/plan")
async def generate_training_plan(
    req: TrainingPlanGenerateRequest = Body(...)
):
    """
    生成光电训练计划接口
    - 入参：TrainingPlanGenerateRequest模型（包含username、training_day_id、training_day_data）
    - 出参：生成的训练计划数据
    """
    # 基础参数校验（模型已做基础校验，此处补充业务校验）
    if not req.training_day_data:
        raise HTTPException(
            status_code=400,
            detail="训练日数据（training_day_data）不能为空"
        )
    
    try:
        # 调用大模型生成训练计划
        plan_data = call_qwen_training_plan(req.training_day_data, req.username)
        
        return {
            "code": 200,
            "message": "训练计划生成成功",
            "plan": plan_data
        }
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"生成训练计划失败：{str(e)}"
        )


# ------------------- 导出训练计划为Word文件接口（仅修改这部分） -------------------
@router.post("/export/plan")
async def export_training_plan(
    req: TrainingPlanExportRequest = Body(...)
):
    """
    导出训练计划为Word文件（返回文件流）
    - 入参：TrainingPlanExportRequest模型（包含username、plan_data）
    - 出参：Word文件二进制流（attachment形式）
    """
    # 基础参数校验
    if not req.plan_data:
        raise HTTPException(
            status_code=400,
            detail="训练计划数据（plan_data）不能为空"
        )
    
    try:
        # 1. 创建Word文档并填充内容
        doc = Document()
        
        # 标题处理（默认标题 + 时间戳，兼容空值）
        plan_title = req.plan_data.get("plan_title", DEFAULT_PLAN_FILENAME)
        safe_title = SAFE_FILENAME_PATTERN.sub("_", plan_title)  # 过滤特殊字符
        
        # 修复点1：移除海象运算符，兼容所有Python版本
        doc.add_heading(plan_title, level=0)
        
        # 添加计划基本信息
        doc.add_heading("计划基本信息", level=1)
        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"生成用户：{req.username}")
        
        # 计划描述（修复空值访问）
        plan_desc = req.plan_data.get("plan_desc", "")
        if plan_desc:
            doc.add_heading("计划描述", level=1)
            doc.add_paragraph(plan_desc)
        
        # 核心目标（修复空值访问）
        core_goal = req.plan_data.get("core_goal", "")
        if core_goal:
            doc.add_heading("核心目标", level=1)
            doc.add_paragraph(core_goal)
        
        # 分天训练计划（修复嵌套空值）
        daily_plans = req.plan_data.get("daily_plans", [])
        if daily_plans:
            doc.add_heading("分天训练计划", level=1)
            for idx, daily_plan in enumerate(daily_plans, 1):
                daily_theme = daily_plan.get("theme", f"第{idx}天训练")
                doc.add_heading(f"第{idx}天：{daily_theme}", level=2)
                
                # 当日目标（修复空值）
                daily_goal = daily_plan.get("daily_goal", "")
                if daily_goal:
                    doc.add_paragraph(f"当日目标：{daily_goal}")
                
                # 项目计划列表（修复嵌套空值）
                project_plans = daily_plan.get("project_plans", [])
                if project_plans:
                    doc.add_heading("项目计划", level=3)
                    for project in project_plans:
                        proj_name = project.get("project_name", "未命名项目")
                        doc.add_paragraph(f"▌项目名称：{proj_name}")
                        
                        # 逐个字段做空值保护
                        training_content = project.get("training_content", "")
                        if training_content:
                            doc.add_paragraph(f"  训练内容：{training_content}")
                        
                        training_methods = project.get("training_methods", [])
                        if training_methods:
                            doc.add_paragraph(f"  训练方法：{', '.join(training_methods)}")
                        
                        acceptance_criteria = project.get("acceptance_criteria", "")
                        if acceptance_criteria:
                            doc.add_paragraph(f"  验收标准：{acceptance_criteria}")
                        
                        notes = project.get("notes", "")
                        if notes:
                            doc.add_paragraph(f"  注意事项：{notes}")
                        
                        doc.add_paragraph("-" * 50)  # 分隔线
                
                # 当日总结（修复空值）
                daily_summary = daily_plan.get("daily_summary", "")
                if daily_summary:
                    doc.add_paragraph(f"当日总结：{daily_summary}")
                
                # 次日预习提示（修复空值）
                next_day_tips = daily_plan.get("next_day_tips", "")
                if next_day_tips:
                    doc.add_paragraph(f"次日预习提示：{next_day_tips}")
                
                doc.add_page_break()  # 分页
        
        # 计划总结（修复空值）
        plan_summary = req.plan_data.get("plan_summary", "")
        if plan_summary:
            doc.add_heading("计划总结", level=1)
            doc.add_paragraph(plan_summary)
        
        # 执行建议（修复空值）
        execution_suggestion = req.plan_data.get("execution_suggestion", "")
        if execution_suggestion:
            doc.add_heading("执行建议", level=1)
            doc.add_paragraph(execution_suggestion)

        # 2. 将文档写入字节流
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)  # 重置文件指针到开头

        # 3. 构造响应（修复中文文件名编码问题）
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        # 修复点2：中文文件名URL编码，避免响应头解析失败
        encoded_filename = urllib.parse.quote(filename)
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                # 修复点3：兼容所有浏览器的文件名格式
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
                "Access-Control-Expose-Headers": "Content-Disposition"  # 允许前端获取文件名
            }
        )
    
    # 修复点4：更详细的异常信息，方便定位问题
    except KeyError as e:
        raise HTTPException(
            status_code=500,
            detail=f"导出训练计划失败：缺失字段 {str(e)}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"导出训练计划失败：{str(e)}，错误类型：{type(e).__name__}"
        )